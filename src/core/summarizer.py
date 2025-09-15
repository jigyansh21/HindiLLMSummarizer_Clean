"""
Core Summarization Service
Handles text processing, AI model integration, and export functionality
"""

import asyncio
import time
import os
import tempfile
from typing import Dict, Any, Literal
from pathlib import Path

# AI Model imports
from transformers.models.t5 import T5Tokenizer, T5ForConditionalGeneration
import torch

# Export libraries
from fpdf import FPDF
from docx import Document
from docx.shared import Inches
import markdown

class SummarizerService:
    def __init__(self):
        self.tokenizer = None
        self.model = None
        self.model_loaded = False
        
    async def load_model(self):
        """Load the T5 model asynchronously"""
        if self.model_loaded:
            return
            
        try:
            print("Loading T5 model...")
            # Run model loading in thread pool to avoid blocking
            loop = asyncio.get_event_loop()
            self.tokenizer, self.model = await loop.run_in_executor(
                None, self._load_model_sync
            )
            self.model_loaded = True
            print("T5 model loaded successfully!")
        except Exception as e:
            print(f"Error loading model: {e}")
            self.model_loaded = False
    
    def _load_model_sync(self):
        """Synchronous model loading"""
        model_name = "t5-small"
        tokenizer = T5Tokenizer.from_pretrained(model_name)
        model = T5ForConditionalGeneration.from_pretrained(model_name)
        return tokenizer, model
    
    async def summarize_text(
        self, 
        text: str, 
        language: Literal["hindi", "english"] = "hindi",
        summary_length: Literal["short", "medium", "long", "auto"] = "auto"
    ) -> Dict[str, Any]:
        """Summarize text using AI model"""
        start_time = time.time()
        
        # Load model if not loaded
        if not self.model_loaded:
            await self.load_model()
        
        # For now, use extractive summarization as primary method
        # T5 model needs more work for proper summarization
        print("Using extractive summarization method")
        return await self._extractive_summarize(text, language, summary_length)
    
    async def _extractive_summarize(
        self, 
        text: str, 
        language: Literal["hindi", "english"],
        summary_length: Literal["short", "medium", "long"]
    ) -> Dict[str, Any]:
        """Improved extractive summarization"""
        start_time = time.time()
        
        # Better sentence splitting for both English and Hindi
        import re
        # Handle both English and Hindi sentence endings
        sentences = re.split(r'[.!?।]+', text)
        sentences = [s.strip() for s in sentences if s.strip() and len(s.strip()) > 10]
        
        word_count = len(text.split())
        
        # Better length mapping based on word count
        if summary_length == "short":
            target_words = min(50, max(20, word_count // 8))
        elif summary_length == "medium":
            target_words = min(100, max(40, word_count // 5))
        else:  # long
            target_words = min(200, max(80, word_count // 3))
        
        # Select sentences to reach target word count
        summary_sentences = []
        current_words = 0
        
        for sentence in sentences:
            sentence_words = len(sentence.split())
            if current_words + sentence_words <= target_words:
                summary_sentences.append(sentence)
                current_words += sentence_words
            else:
                # Add partial sentence if we're close to target
                if current_words < target_words * 0.7:  # If we're less than 70% of target
                    remaining_words = target_words - current_words
                    words = sentence.split()
                    if len(words) > remaining_words:
                        partial_sentence = ' '.join(words[:remaining_words])
                        summary_sentences.append(partial_sentence + "...")
                    else:
                        summary_sentences.append(sentence)
                break
        
        # If no sentences selected, take first sentence
        if not summary_sentences:
            summary_sentences = [sentences[0]] if sentences else [text[:100] + "..."]
        
        summary = '. '.join(summary_sentences).strip()
        
        # Ensure summary ends with proper punctuation
        if summary and not summary.endswith(('.', '!', '?')):
            summary += '.'
        
        processing_time = time.time() - start_time
        
        return {
            "summary": summary,
            "original_length": word_count,
            "summary_length": len(summary.split()),
            "compression_ratio": round(len(summary.split()) / word_count, 2),
            "processing_time": round(processing_time, 2)
        }
    
    async def summarize_url(
        self, 
        url: str, 
        language: Literal["hindi", "english"] = "hindi",
        summary_length: Literal["short", "medium", "long", "auto"] = "auto"
    ) -> Dict[str, Any]:
        """Extract and summarize content from URL"""
        try:
            from newspaper import Article
            
            # Extract article content
            article = Article(url, language="hi" if language == "hindi" else "en")
            article.download()
            article.parse()
            
            text = article.text.strip()
            title = article.title or "Untitled Article"
            
            if not text:
                raise ValueError("No content extracted from URL")
            
            # Summarize the extracted text
            result = await self.summarize_text(text, language, summary_length)
            result["title"] = title
            
            return result
            
        except Exception as e:
            raise Exception(f"Failed to process URL: {str(e)}")
    
    async def export_pdf(
        self, 
        summary: str, 
        title: str = "Summary",
        language: Literal["hindi", "english"] = "hindi"
    ) -> str:
        """Export summary as PDF with proper font support"""
        try:
            pdf = FPDF()
            pdf.add_page()
            
            # Set up fonts
            font_path = os.path.join(os.path.dirname(__file__), "..", "..", "fonts", "NotoSansDevanagari-Regular.ttf")
            if os.path.exists(font_path):
                pdf.add_font("NotoSans", "", font_path, uni=True)
                pdf.set_font("NotoSans", "", 12)
            else:
                pdf.set_font("Arial", "", 12)
            
            # Add title
            pdf.set_font("Arial", "B", 16)
            pdf.cell(0, 10, title, 0, 1, "C")
            pdf.ln(10)
            
            # Add summary
            pdf.set_font("NotoSans" if os.path.exists(font_path) else "Arial", "", 12)
            pdf.multi_cell(0, 8, summary)
            
            # Add footer
            pdf.ln(20)
            pdf.set_font("Arial", "I", 8)
            pdf.cell(0, 10, f"Generated by MultiLanguage AI Text Summarizer - {time.strftime('%Y-%m-%d %H:%M')}", 0, 1, "C")
            
            # Save to temporary file
            temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=".pdf")
            pdf.output(temp_file.name)
            
            return temp_file.name
            
        except Exception as e:
            raise Exception(f"Failed to create PDF: {str(e)}")
    
    async def export_word(
        self, 
        summary: str, 
        title: str = "Summary",
        language: Literal["hindi", "english"] = "hindi"
    ) -> str:
        """Export summary as Word document"""
        try:
            doc = Document()
            
            # Add title
            title_para = doc.add_heading(title, 0)
            
            # Add summary
            summary_para = doc.add_paragraph(summary)
            
            # Add metadata
            doc.add_paragraph(f"\n\nGenerated by MultiLanguage AI Text Summarizer")
            doc.add_paragraph(f"Date: {time.strftime('%Y-%m-%d %H:%M')}")
            
            # Save to temporary file
            temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
            doc.save(temp_file.name)
            
            return temp_file.name
            
        except Exception as e:
            raise Exception(f"Failed to create Word document: {str(e)}")
    
    async def export_markdown(
        self, 
        summary: str, 
        title: str = "Summary",
        language: Literal["hindi", "english"] = "hindi"
    ) -> str:
        """Export summary as Markdown"""
        try:
            markdown_content = f"""# {title}

{summary}

---

*Generated by MultiLanguage AI Text Summarizer*  
*Date: {time.strftime('%Y-%m-%d %H:%M')}*
"""
            
            # Save to temporary file
            temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=".md", mode='w', encoding='utf-8')
            temp_file.write(markdown_content)
            temp_file.close()
            
            return temp_file.name
            
        except Exception as e:
            raise Exception(f"Failed to create Markdown file: {str(e)}")
